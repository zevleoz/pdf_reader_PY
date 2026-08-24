"""极简 Web 前端：4 份 PDF (A2/B3/B4/B6) → 生成综合 PDF → 下载。

路由：
  GET  /                      → 上传页面（4 个带命名的上传槽 + 底部生成按钮）
  POST /api/generate          → 接收 4 份 PDF，顺序执行 extract → validate → generate，
                                 然后把 output/report.pdf 返回为附件
  GET  /output/<filename>     → 直接下载生成的文件
"""

from __future__ import annotations

import json
import os
import sys
import traceback
from datetime import date, timedelta
from functools import wraps
from pathlib import Path
from typing import List, Optional

from flask import (Flask, jsonify, render_template, request,
                   send_from_directory, abort, session, redirect)

import extract
import validate
import generate as _generate_module
from data_points import apply_report_data
import db as _db

BASE_DIR = Path(__file__).resolve().parent
INPUT_DIR = BASE_DIR / "input"
OUTPUT_DIR = BASE_DIR / "output"
DATA_DIR = BASE_DIR / "data"
TEMPLATE_DIR = BASE_DIR / "templates"
BRANDING_DIR = BASE_DIR / "branding"

INPUT_DIR.mkdir(parents=True, exist_ok=True)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
DATA_DIR.mkdir(parents=True, exist_ok=True)
_db.init_db()

REQUIRED_KEYS: List[str] = ["A2", "B3", "B4", "B6"]

app = Flask(__name__, template_folder=str(TEMPLATE_DIR),
            static_folder=str(OUTPUT_DIR), static_url_path="/output")

app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024
app.secret_key = os.environ.get('SECRET_KEY', 'dev-key-change-in-production')

ADMIN_PASSWORD = os.environ.get('ADMIN_PASSWORD', 'y4admin2026')


def admin_required(f):
    """Decorator: require admin session for mutating endpoints."""
    @wraps(f)
    def wrapper(*args, **kwargs):
        if not session.get('is_admin'):
            return jsonify({"ok": False, "error": "需要管理员权限"}), 403
        return f(*args, **kwargs)
    return wrapper


def page_login_required(f):
    """Decorator: require admin session for page routes. Redirects to /login if not logged in."""
    @wraps(f)
    def wrapper(*args, **kwargs):
        if not session.get('is_admin'):
            return redirect('/login?next=' + request.path)
        return f(*args, **kwargs)
    return wrapper


@app.route("/style.css")
def serve_style():
    """Serve shared design-system CSS from templates/style.css."""
    css_path = TEMPLATE_DIR / "style.css"
    return send_from_directory(str(TEMPLATE_DIR), "style.css", mimetype="text/css")


@app.route("/branding/<path:filename>")
def branding(filename):
    """Serve branding assets (logo, watermark) directly from the branding/ folder.

    Needed so the topbar logo can display before generate.py has ever run
    (generate.py copies branding → output/branding/ only once a report is produced).
    """
    target = BRANDING_DIR / filename
    if not target.exists():
        abort(404)
    return send_from_directory(str(BRANDING_DIR), filename, as_attachment=False)


@app.errorhandler(413)
def request_entity_too_large(error):
    return jsonify({"ok": False, "error": "文件太大，单文件不超过 50MB"}), 413


@app.route("/")
def landing():
    return render_template("landing.html")


@app.route("/login")
def login_page():
    return render_template("login.html")


@app.route("/generate")
@page_login_required
def index():
    return render_template("index.html")


@app.route("/preview")
@page_login_required
def preview():
    from generate import build_view_data, render_html
    from data_points import apply_report_data
    apply_report_data()
    view_data = build_view_data()
    from pathlib import Path as P
    output_path = OUTPUT_DIR / "preview.html"
    render_html(view_data, output_path)
    return send_from_directory(str(OUTPUT_DIR), "preview.html")


# ---------------------------------------------------------------------------
# 进度查询接口（供前端进度条轮询）
# ---------------------------------------------------------------------------
@app.route("/api/progress", methods=["GET"])
def api_progress():
    prog_file = DATA_DIR / "_progress.json"
    if prog_file.exists():
        try:
            return jsonify(json.loads(prog_file.read_text(encoding="utf-8")))
        except Exception:
            pass
    return jsonify({"stage": "idle", "percent": 0, "message": ""})


# ---------------------------------------------------------------------------
# 核心接口：接收 4 份 PDF → 跑管道 → 返回 report.pdf
# ---------------------------------------------------------------------------
@app.route("/api/generate", methods=["POST"])
def api_generate():
    # 1) 接收1-4个文件，不再强制要求4个
    files_by_key = {}
    for key in REQUIRED_KEYS:
        f = request.files.get(key)
        if f and f.filename:
            files_by_key[key] = f
    
    if not files_by_key:
        return jsonify({
            "ok": False,
            "error": "请至少上传一份PDF文件"
        }), 400

    # 2) 清空 input/ 旧文件，按 report_<KEY>.pdf 保存
    try:
        for old in INPUT_DIR.glob("*.pdf"):
            old.unlink()
    except OSError:
        pass

    saved_names = {}
    try:
        for key, f in files_by_key.items():
            # 保留原始文件名中的版本信息（初中/高中）
            original_name = f.filename or ""
            if key == "B6" and ("高中" in original_name or "初中" in original_name):
                # 提取版本关键词，附加到标准文件名
                version = "高中" if "高中" in original_name else "初中"
                target = INPUT_DIR / f"report_{key}_{version}.pdf"
            else:
                target = INPUT_DIR / f"report_{key}.pdf"
            f.save(str(target))
            saved_names[key] = f.filename
    except Exception as exc:
        return jsonify({"ok": False,
                         "error": f"保存上传文件失败: {exc}"}), 500

    try:
        # 3) extract：解析 4 份 PDF → data/report_data.json
        #    ⚠️  extract.py 现在强制依赖视觉 OCR API；如果 API 未配置或
        #    调用失败，extract.main() 会抛出 RuntimeError，此处直接转
        #    成 JSON 错误响应给前端（浏览器弹窗提示）。
        rc = extract.main()
        if rc != 0:
            return jsonify({
                "ok": False,
                "error": f"提取数据失败 (extract.main() 返回 {rc})。请检查输入 PDF 是否可读，或查看日志。"
            }), 500

        # 4) validate：做一次完整性校验（非阻塞，失败也继续）
        try:
            validate.main()
        except Exception:
            pass

        # 5) 关键点：让 data_points 基于新的 report_data.json 重新填充 USER_DATA
        apply_result = apply_report_data()

        # 5b) 如果用户手动输入了思维模式分值，覆盖提取的数据
        mindset_score = request.form.get('mindset_score')
        if mindset_score:
            try:
                score_val = float(mindset_score)
                if 0 <= score_val <= 100:
                    from data_points import USER_DATA
                    USER_DATA['059'] = str(score_val)
                    print(f"[思维模式] 用户手动输入分值: {score_val}")

                    # 同时更新 report_data.json，避免被 build_view_data() 中的
                    # apply_report_data() 覆盖
                    report_data_path = DATA_DIR / "report_data.json"
                    if report_data_path.exists():
                        try:
                            with open(report_data_path, 'r', encoding='utf-8') as f:
                                report_data = json.load(f)
                            for item in report_data.get('schema_124', []):
                                if item.get('code') == '059':
                                    item['value'] = str(score_val)
                                    break
                            with open(report_data_path, 'w', encoding='utf-8') as f:
                                json.dump(report_data, f, ensure_ascii=False, indent=2)
                            print(f"[思维模式] 已更新 report_data.json 中的 059 值为 {score_val}")
                        except Exception as e:
                            print(f"[思维模式] 更新 report_data.json 失败: {e}")
            except ValueError:
                pass

        # 6) 清理 output/ 旧产物，避免 chrome 基于旧文件命名出错
        #    清理所有 report.pdf/html 以及动态命名的 凭远Y4评测报告_*.pdf/html
        for pattern in ["report.pdf", "report.html", "凭远Y4评测报告_*.pdf", "凭远Y4评测报告_*.html"]:
            for old in OUTPUT_DIR.glob(pattern):
                try:
                    old.unlink()
                except OSError:
                    pass

        # 7) 重新 import generate 模块，使 build_view_data 里读取 USER_DATA 的值是最新的
        #    注：generate 的函数/常量会引用 data_points.USER_DATA（全局），
        #    在 apply_report_data 之后，值已经被更新。不需要 reload，直接跑 main 即可。
        try:
            _generate_module.main()
        except SystemExit as exc:
            if exc.code not in (None, 0):
                return jsonify({"ok": False,
                                 "error": f"生成 PDF 失败: SystemExit({exc.code})"}), 500
        except Exception as exc:
            tb = traceback.format_exc()
            return jsonify({"ok": False,
                             "error": f"生成 PDF 时异常: {exc}",
                             "trace": tb}), 500

        # Find the generated PDF (could be named with student name or report.pdf)
        pdf_files = sorted(OUTPUT_DIR.glob("*.pdf"))
        pdf_path = None
        if pdf_files:
            # Prefer the dynamically named file
            named = [f for f in pdf_files if "凭远Y4评测报告" in f.name]
            if named:
                pdf_path = named[0]
            else:
                pdf_path = pdf_files[0]

        if not pdf_path or not pdf_path.exists():
            import subprocess as _sp
            chrome_found = None
            for p in ["/usr/bin/chromium-browser", "/usr/bin/chromium",
                      "/usr/bin/google-chrome", "/usr/bin/google-chrome-stable"]:
                if Path(p).exists():
                    chrome_found = p
                    break
            if not chrome_found:
                for cmd in ["chromium-browser", "chromium", "google-chrome", "google-chrome-stable"]:
                    try:
                        r = _sp.run(["which", cmd], capture_output=True, text=True)
                        if r.returncode == 0 and r.stdout.strip():
                            chrome_found = r.stdout.strip()
                            break
                    except Exception:
                        pass
            return jsonify({
                "ok": False,
                "error": f"生成流程完成，但未在 output/ 下找到 PDF。Chrome 检测: {chrome_found or '未找到'}。请检查 Chrome 是否可用。",
                "chrome_path": chrome_found,
            }), 500

        # 8) Save to database (non-blocking, best-effort)
        try:
            report_data_path = DATA_DIR / "report_data.json"
            if report_data_path.exists():
                rd = json.loads(report_data_path.read_text(encoding="utf-8"))
                student_info = rd.get("student", {}) or {}
                sname = student_info.get("name", "").strip()

                # Use student_id from dropdown if provided, else manual name, else from PDF
                selected_sid = request.form.get("student_id", "").strip()
                manual_name = request.form.get("manual_student_name", "").strip()

                if selected_sid:
                    # Use the pre-selected student from booking
                    sid = int(selected_sid)
                    sname = sname or manual_name or f"学生#{sid}"
                    # Update student info from PDF if available
                    if student_info:
                        _db.update_student(sid, **{k: v for k, v in {
                            "gender": student_info.get("gender", ""),
                            "birthday": student_info.get("birthday", ""),
                            "grade": student_info.get("grade", ""),
                        }.items() if v})
                    print(f"[DB] 关联到预约学生: {sname} (id={sid})")
                elif manual_name:
                    sid = _db.find_or_create_student(name=manual_name)
                    sname = manual_name
                    print(f"[DB] 手动输入学生: {sname} (id={sid})")
                elif sname:
                    sid = _db.find_or_create_student(
                        name=sname,
                        gender=student_info.get("gender", ""),
                        birthday=student_info.get("birthday", ""),
                        grade=student_info.get("grade", ""),
                    )
                    print(f"[DB] 从PDF提取学生: {sname} (id={sid})")
                else:
                    sid = None

                if sid:
                    _db.add_report(
                        student_id=sid,
                        report_date=date.today(),
                        pdf_path=str(pdf_path),
                        data_json=report_data_path.read_text(encoding="utf-8"),
                    )
        except Exception as e:
            print(f"[DB] 保存报告失败 (非致命): {e}")

        # 9) 返回 PDF 作为附件
        download_filename = pdf_path.name
        resp = send_from_directory(
            str(OUTPUT_DIR), pdf_path.name,
            as_attachment=True,
            download_name=download_filename,
            mimetype="application/pdf",
        )
        resp.headers["X-Applied-Items"] = str(apply_result.get("applied", 0))
        resp.headers["X-Total-Items"] = str(apply_result.get("total_items", 0))
        return resp

    except Exception as exc:
        tb = traceback.format_exc()
        return jsonify({"ok": False,
                         "error": f"服务端异常: {exc}",
                         "trace": tb}), 500


# ---------------------------------------------------------------------------
# AI 聊天接口：接收用户消息 + 历史对话 → 调 DashScope 文本 LLM → 返回回复
# ---------------------------------------------------------------------------
@app.route("/api/chat", methods=["POST"])
def api_chat():
    import urllib.request as _ureq

    data = request.get_json(force=True)
    user_message = data.get("message", "").strip()
    history = data.get("history", [])

    if not user_message and not history:
        return jsonify({"ok": False, "error": "消息不能为空"}), 400

    # 1) 读取 AI prompt
    prompt_path = BASE_DIR / "prompts" / "ai_interpreter.md"
    system_prompt = prompt_path.read_text(encoding="utf-8") if prompt_path.exists() else "你是测评解读助手。"

    # 2) 读取 report_data.json 作为上下文
    report_path = DATA_DIR / "report_data.json"
    if report_path.exists():
        report_data = json.loads(report_path.read_text(encoding="utf-8"))
        schema_items = report_data.get("schema_124", [])
        data_text = "\n".join(
            f"{it.get('code','?')} {it.get('label','?')}：{it.get('value', '—')}"
            for it in schema_items if it.get("value")
        )
        student = report_data.get("student", {})
        student_text = f"学生：{student.get('name','—')}，{student.get('gender','—')}，{student.get('grade','—')}"
        context = f"{student_text}\n\n测评数据：\n{data_text}"
    else:
        context = "（暂无测评数据）"

    # 3) 组装 messages
    messages = [
        {"role": "system", "content": system_prompt + "\n\n以下是学生测评数据：\n" + context},
    ]
    messages.extend(history[-10:])
    if user_message:
        messages.append({"role": "user", "content": user_message})

    # 4) 调用 DashScope OpenAI 兼容接口
    dashscope_key = os.environ.get("DASHSCOPE_API_KEY", extract.DEFAULT_DASHSCOPE_KEY).strip()
    url = "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions"
    payload = json.dumps({
        "model": os.environ.get("AI_TEXT_MODEL", "qwen-plus"),
        "messages": messages,
        "temperature": 0.5,
        "max_tokens": 8192,
    }).encode("utf-8")
    req = _ureq.Request(
        url, data=payload,
        headers={"Authorization": f"Bearer {dashscope_key}",
                 "Content-Type": "application/json"},
        method="POST",
    )
    try:
        with _ureq.urlopen(req, timeout=120) as resp:
            result = json.loads(resp.read().decode("utf-8"))
        reply = result["choices"][0]["message"]["content"]
        return jsonify({"ok": True, "reply": reply})
    except Exception as exc:
        return jsonify({"ok": False, "error": f"AI 调用失败: {exc}"}), 500


# ---------------------------------------------------------------------------
# 静态输出文件访问
# ---------------------------------------------------------------------------
@app.route("/output/<path:filename>")
def download(filename):
    target = OUTPUT_DIR / filename
    if not target.exists():
        abort(404)
    return send_from_directory(str(OUTPUT_DIR), filename, as_attachment=False)


# ---------------------------------------------------------------------------
# 解读会会议纪要生成
# ---------------------------------------------------------------------------
@app.route("/transcript")
@page_login_required
def transcript_page():
    """Render the transcript upload / summary generation page."""
    return render_template("transcript.html")


@app.route("/api/transcript", methods=["POST"])
def api_transcript():
    import urllib.request as _ureq

    data = request.get_json(force=True)
    transcript_text = (data.get("transcript") or "").strip()
    student_name = (data.get("student_name") or "").strip()
    student_grade = (data.get("student_grade") or "").strip()
    student_gender = (data.get("student_gender") or "").strip()

    if len(transcript_text) < 50:
        return jsonify({"ok": False, "error": "逐字稿太短，至少 50 字"}), 400

    # 1) 读取 transcript prompt
    prompt_path = BASE_DIR / "prompts" / "transcript_summary.md"
    system_prompt = prompt_path.read_text(encoding="utf-8") if prompt_path.exists() else "你是解读会纪要撰写人。"

    # 2) 读取 report_data.json 作为上下文（如果存在）
    report_summary = "（暂无测评数据）"
    report_path = DATA_DIR / "report_data.json"
    if report_path.exists():
        try:
            rd = json.loads(report_path.read_text(encoding="utf-8"))
            student_info = rd.get("student", {})
            s124 = rd.get("schema_124", [])
            lines = []
            for item in s124:
                val = item.get("value", "")
                if val and val not in ("", "—", None):
                    lines.append(f"  {item.get('label', item.get('code', '?'))}: {val}")
            if lines:
                report_summary = "\n".join(lines[:60])
        except Exception:
            pass

    # 3) 组装 prompt 变量
    system_prompt = system_prompt.replace("{student_name}", student_name or "未填写")
    system_prompt = system_prompt.replace("{report_summary}", report_summary)
    system_prompt = system_prompt.replace("{transcript}", transcript_text[:12000])

    # 4) 组装 messages
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": "请根据以上逐字稿和测评数据，撰写完整的解读会会议纪要。"},
    ]

    # 5) 调用 DashScope API
    dashscope_key = os.environ.get("DASHSCOPE_API_KEY", extract.DEFAULT_DASHSCOPE_KEY).strip()
    url = "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions"
    payload = json.dumps({
        "model": os.environ.get("AI_TEXT_MODEL", "qwen-plus"),
        "messages": messages,
        "temperature": 0.7,
        "max_tokens": 8192,
    }).encode("utf-8")
    req = _ureq.Request(
        url, data=payload,
        headers={"Authorization": f"Bearer {dashscope_key}",
                 "Content-Type": "application/json"},
        method="POST",
    )
    try:
        with _ureq.urlopen(req, timeout=180) as resp:
            result = json.loads(resp.read().decode("utf-8"))
        summary = result["choices"][0]["message"]["content"]
        return jsonify({"ok": True, "summary": summary})
    except Exception as exc:
        return jsonify({"ok": False, "error": f"AI 调用失败: {exc}"}), 500


def _build_docx(summary_text: str, student_name: str = "", student_grade: str = "",
               student_gender: str = "") -> bytes:
    """Convert markdown-style meeting minutes text into a styled Word .docx file.
    Returns the raw bytes of the docx file.
    """
    import re
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ---- Page margins ----
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ---- Base styles ----
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x14, 0x14, 0x14)

    # ---- Title ----
    title_line = "凭远教育 · Y4 解读会会议纪要"
    if student_name:
        title_line += f" — {student_name}"
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(title_line)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.color.rgb = RGBColor(0x14, 0x14, 0x14)
    title_run.font.name = "微软雅黑"
    title_p.paragraph_format.space_after = Pt(6)

    # ---- Student info table ----
    from datetime import datetime
    info_line = f"学生：{student_name or '—'}"
    if student_grade:
        info_line += f"　｜　年级：{student_grade}"
    if student_gender:
        info_line += f"　｜　性别：{student_gender}"
    info_line += f"　｜　生成时间：{datetime.now().strftime('%Y-%m-%d')}"
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run(info_line)
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    meta_p.paragraph_format.space_after = Pt(16)

    # ---- Red divider ----
    div_p = doc.add_paragraph()
    div_run = div_p.add_run("━" * 30)
    div_run.font.color.rgb = RGBColor(0xB3, 0x3A, 0x3A)
    div_run.font.size = Pt(10)
    div_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_p.paragraph_format.space_after = Pt(12)

    # ---- Parse lines and build document structure ----
    # Heading patterns: **一、** **二、** etc
    heading_re = re.compile(r"^\s*\*\*(.+?)\*\*\s*(?:（(.+?)）)?\s*$")
    # Bullet: starts with • or - * or • [ ]
    bullet_chars = ("•", "·", "-", "*", "【")

    lines = summary_text.split("\n")
    section_title = None
    section_subtitle = None

    for raw in lines:
        line = raw.rstrip()
        stripped = line.strip()

        # Skip separator lines like "---"
        if re.match(r"^-+$", stripped) or re.match(r"^=+$", stripped) or re.match(r"^━+$", stripped):
            continue
        if stripped.startswith("> ") or stripped == ">":
            # Markdown quote line → treat as italic small paragraph
            if stripped == ">":
                continue
            q = stripped[2:].strip()
            if q:
                p = doc.add_paragraph()
                run = p.add_run(q)
                run.italic = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)
            continue

        # Heading check: **一、核心发现 · 3 条**（...）
        m = heading_re.match(stripped)
        if m:
            section_title = m.group(1).strip()
            section_subtitle = m.group(2).strip() if m.group(2) else None
            hp = doc.add_paragraph()
            hrun = hp.add_run(section_title)
            hrun.bold = True
            hrun.font.size = Pt(14)
            hrun.font.color.rgb = RGBColor(0xB3, 0x3A, 0x3A)
            hrun.font.name = "微软雅黑"
            hp.paragraph_format.space_before = Pt(12)
            hp.paragraph_format.space_after = Pt(4)
            if section_subtitle:
                subrun = hp.add_run(f" （{section_subtitle}）")
                subrun.bold = False
                subrun.font.size = Pt(10)
                subrun.font.color.rgb = RGBColor(0x8A, 0x8A, 0x8A)
            continue

        if not stripped:
            continue  # empty line

        # Check if it's a sub-heading line (bold **...** without a bullet)
        if stripped.startswith("**") and stripped.endswith("**") and not any(stripped.startswith(b) for b in bullet_chars):
            subh = stripped[2:-2]
            hp2 = doc.add_paragraph()
            r2 = hp2.add_run(subh)
            r2.bold = True
            r2.font.size = Pt(11.5)
            r2.font.color.rgb = RGBColor(0x14, 0x14, 0x14)
            hp2.paragraph_format.space_before = Pt(8)
            hp2.paragraph_format.space_after = Pt(2)
            continue

        # Bullet lines (starts with bullet char, possibly after whitespace)
        is_bullet = False
        content = stripped
        checkbox = ""
        # Strip leading bullet marker
        lead = re.match(r"^(\s*[-•·*·]\s*(?:\[\s*[xX\s]\]\s*)?)", content)
        if lead:
            marker = lead.group(1)
            if "[ ]" in marker or "［］" in marker or "[  ]" in marker:
                checkbox = "☐"
            elif re.search(r"\[\s*[xX]\s*\]", marker):
                checkbox = "☑"
            content = content[lead.end():].strip()
            is_bullet = True

        if not is_bullet and (content.startswith("【") or content.startswith("• ") or content.startswith("· ")):
            is_bullet = True
            if content.startswith("• "):
                content = content[2:]
            elif content.startswith("· "):
                content = content[2:]

        # Strip inline **...** bold markers and convert to rich runs
        tokens = re.split(r"(\*\*[^*]+\*\*)", content)
        if is_bullet:
            bullet_prefix = f"{checkbox} • " if checkbox else "• "
            p = doc.add_paragraph(style="List Bullet")
            # docx default bullet style might look odd; just prepend char
            # Override: re-add as normal paragraph with bullet marker
            try:
                p.clear()
            except Exception:
                pass
            # Build content without list style to avoid indent inconsistencies
            p2 = doc.add_paragraph()
            b_run = p2.add_run(bullet_prefix)
            b_run.font.color.rgb = RGBColor(0xB3, 0x3A, 0x3A)
            b_run.bold = True
            p2.paragraph_format.space_before = Pt(2)
            p2.paragraph_format.space_after = Pt(2)
            p2.paragraph_format.left_indent = Cm(0.6)
            # Remove the duplicate bullet paragraph 'p'
            try:
                p_elm = p._element
                p_elm.getparent().remove(p_elm)
            except Exception:
                pass

            # Append content tokens
            for tok in tokens:
                if tok.startswith("**") and tok.endswith("**"):
                    r = p2.add_run(tok[2:-2])
                    r.bold = True
                    r.font.size = Pt(11)
                else:
                    r = p2.add_run(tok)
                    r.font.size = Pt(11)
            continue

        # Plain paragraph (non-heading, non-bullet)
        pp = doc.add_paragraph()
        for tok in tokens:
            if tok.startswith("**") and tok.endswith("**"):
                r = pp.add_run(tok[2:-2])
                r.bold = True
                r.font.size = Pt(11)
            else:
                r = pp.add_run(tok)
                r.font.size = Pt(11)

    # ---- Footer divider ----
    doc.add_paragraph()
    footer_div = doc.add_paragraph()
    fd_run = footer_div.add_run("— 凭远教育 · Y4 综合测评系统 —")
    fd_run.font.size = Pt(9)
    fd_run.font.color.rgb = RGBColor(0x8A, 0x8A, 0x8A)
    footer_div.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@app.route("/api/transcript/docx", methods=["POST"])
def api_transcript_docx():
    """Generate minutes as DOCX file. Accepts JSON with summary + student info,
    returns the .docx binary as attachment download.
    """
    data = request.get_json(force=True)
    summary_text = (data.get("summary") or "").strip()
    student_name = (data.get("student_name") or "").strip()
    student_grade = (data.get("student_grade") or "").strip()
    student_gender = (data.get("student_gender") or "").strip()

    if not summary_text:
        return jsonify({"ok": False, "error": "纪要内容为空"}), 400

    try:
        docx_bytes = _build_docx(summary_text, student_name, student_grade, student_gender)
    except Exception as exc:
        import traceback
        tb = traceback.format_exc()
        return jsonify({"ok": False, "error": f"DOCX 生成失败: {exc}", "trace": tb}), 500

    filename_suffix = student_name if student_name else "纪要"
    filename = f"凭远Y4解读会纪要_{filename_suffix}.docx"
    safe_name = filename.encode("utf-8").decode("latin-1", "ignore")
    from flask import Response
    resp = Response(
        docx_bytes,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"attachment; filename=\"{safe_name}\"; filename*=UTF-8''{filename}"
        }
    )
    return resp


# ---------------------------------------------------------------------------
# Student management
# ---------------------------------------------------------------------------
@app.route("/students")
@page_login_required
def students_page():
    return render_template("students.html")


@app.route("/api/students")
def api_students():
    students = _db.get_students()
    return jsonify({"ok": True, "students": students})


@app.route("/api/students/<int:student_id>/reports")
def api_student_reports(student_id):
    reports = _db.get_student_reports(student_id)
    return jsonify({"ok": True, "reports": reports})


# ---------------------------------------------------------------------------
# Admin authentication
# ---------------------------------------------------------------------------
@app.route("/api/admin/login", methods=["POST"])
def api_admin_login():
    data = request.get_json(force=True)
    password = data.get("password", "")
    if password == ADMIN_PASSWORD:
        session['is_admin'] = True
        return jsonify({"ok": True})
    return jsonify({"ok": False, "error": "密码错误"}), 401


@app.route("/api/admin/logout", methods=["POST"])
def api_admin_logout():
    session.pop('is_admin', None)
    return jsonify({"ok": True})


@app.route("/api/admin/check")
def api_admin_check():
    return jsonify({"ok": True, "is_admin": bool(session.get('is_admin'))})


# ---------------------------------------------------------------------------
# Booking system
# ---------------------------------------------------------------------------
@app.route("/booking")
def booking_page():
    return render_template("booking.html")


@app.route("/api/booking", methods=["POST"])
def api_booking():
    from datetime import datetime
    data = request.get_json(force=True)
    name = (data.get("student_name") or "").strip()
    if not name:
        return jsonify({"ok": False, "error": "学生姓名必填"}), 400
    appt_time_str = data.get("appointment_time", "")
    try:
        appt_time = datetime.fromisoformat(appt_time_str)
    except ValueError:
        return jsonify({"ok": False, "error": "时间格式错误"}), 400

    # Check slot capacity
    booking_date = appt_time.date()
    time_slot_str = appt_time.strftime("%H:%M")
    booking_counts = _db.get_slot_booking_counts(booking_date)
    booked = booking_counts.get(time_slot_str, 0)
    if booked >= SLOT_CAPACITY:
        return jsonify({"ok": False, "error": "该时段已满（" + str(booked) + "/" + str(SLOT_CAPACITY) + "），请选择其他时间"}), 409

    # Check admin hasn't closed this slot
    avail_entries = _db.get_availability(booking_date)
    avail_map = {e["time_slot"]: e["is_available"] for e in avail_entries}
    is_open = avail_map.get(time_slot_str, True)
    if not is_open:
        return jsonify({"ok": False, "error": "该时段已关闭，请选择其他时间"}), 409

    # Create student + booking in one transaction (auto-archive)
    student_id, booking_id = _db.create_booking_with_student(
        student_name=name,
        appointment_time=appt_time,
        advisor_name=(data.get("advisor_name") or "").strip(),
        school=(data.get("school") or "").strip(),
        single_parent=data.get("single_parent", "false"),
        notes=data.get("notes", ""),
    )
    return jsonify({"ok": True, "booking_id": booking_id, "student_id": student_id})


@app.route("/admin/bookings")
@page_login_required
def admin_bookings_page():
    return render_template("admin_bookings.html")


@app.route("/api/bookings")
def api_bookings():
    status = request.args.get("status")
    bookings = _db.get_bookings(status=status)
    return jsonify({"ok": True, "bookings": bookings})


@app.route("/api/booking/<int:booking_id>/complete", methods=["POST"])
def api_booking_complete(booking_id):
    try:
        sid = _db.complete_booking(booking_id)
        return jsonify({"ok": True, "student_id": sid})
    except ValueError as e:
        return jsonify({"ok": False, "error": str(e)}), 404


@app.route("/api/booking/<int:booking_id>/cancel", methods=["POST"])
def api_booking_cancel(booking_id):
    _db.update_booking_status(booking_id, "cancelled")
    return jsonify({"ok": True})


# ---------------------------------------------------------------------------
# Availability management
# ---------------------------------------------------------------------------
SLOT_CAPACITY = 4  # Max students per time slot


@app.route("/api/availability")
def api_get_availability():
    """Get availability for a given date. Query param: date=YYYY-MM-DD
    Returns each slot with is_available, booked_count, and capacity.
    Default: all slots available. Admin can close slots. Full slots (4/4) are unavailable.
    """
    date_str = request.args.get("date", "")
    if not date_str:
        return jsonify({"ok": False, "error": "请指定日期"}), 400
    try:
        date_val = date.fromisoformat(date_str)
    except ValueError:
        return jsonify({"ok": False, "error": "日期格式错误"}), 400
    entries = _db.get_availability(date_val)
    booking_counts = _db.get_slot_booking_counts(date_val)
    # Build full slot list: no record = available (admin closes slots manually)
    all_slots = []
    available_map = {e["time_slot"]: e["is_available"] for e in entries}
    for ts in _db.TIME_SLOTS:
        is_open = available_map.get(ts, True)  # Default: available
        booked = booking_counts.get(ts, 0)
        is_av = is_open and booked < SLOT_CAPACITY
        all_slots.append({
            "time_slot": ts,
            "is_available": is_av,
            "is_open": is_open,
            "booked_count": booked,
            "capacity": SLOT_CAPACITY,
        })
    return jsonify({"ok": True, "date": date_str, "slots": all_slots})


@app.route("/api/availability/month")
def api_availability_month():
    """Get availability + booking counts for next 30 days (for admin calendar management)."""
    today = date.today()
    end = today + timedelta(days=29)
    range_data = _db.get_availability_range(today, end)
    booking_counts = _db.get_booking_counts_range(today, end)
    return jsonify({
        "ok": True,
        "start": today.isoformat(),
        "end": end.isoformat(),
        "data": range_data,
        "bookings": booking_counts,
    })


@app.route("/api/availability", methods=["POST"])
@admin_required
def api_set_availability():
    """Batch set availability for a date. Admin only.
    Body: {"date": "YYYY-MM-DD", "slots": [{"time_slot": "09:00", "is_available": true}, ...]}
    """
    data = request.get_json(force=True)
    date_str = data.get("date", "")
    slots = data.get("slots", [])
    if not date_str:
        return jsonify({"ok": False, "error": "请指定日期"}), 400
    try:
        date_val = date.fromisoformat(date_str)
    except ValueError:
        return jsonify({"ok": False, "error": "日期格式错误"}), 400
    _db.batch_set_availability(date_val, slots)
    return jsonify({"ok": True})


# ---------------------------------------------------------------------------
# Delete operations
# ---------------------------------------------------------------------------
@app.route("/api/students/<int:student_id>", methods=["DELETE"])
def api_delete_student(student_id):
    try:
        _db.delete_student(student_id)
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/reports/<int:report_id>", methods=["DELETE"])
def api_delete_report(report_id):
    try:
        _db.delete_report(report_id)
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/booking/<int:booking_id>", methods=["DELETE"])
def api_delete_booking(booking_id):
    try:
        _db.delete_booking(booking_id)
        return jsonify({"ok": True})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


# ---------------------------------------------------------------------------
# CSV export
# ---------------------------------------------------------------------------
@app.route("/api/export")
def api_export():
    import csv
    import io
    reports = _db.get_all_reports()
    if not reports:
        return jsonify({"ok": False, "error": "暂无数据"}), 404

    # Collect all field codes across all reports
    all_codes = set()
    for r in reports:
        all_codes.update(r["data"].keys())
    sorted_codes = sorted(all_codes)

    output = io.StringIO()
    writer = csv.writer(output)
    header = ["报告ID", "学生姓名", "性别", "年级", "报告日期", "PDF路径"] + sorted_codes
    writer.writerow(header)
    for r in reports:
        row = [
            r["report_id"], r["student_name"], r["gender"] or "",
            r["grade"] or "", r["report_date"] or "", r["pdf_path"] or "",
        ]
        row.extend(r["data"].get(c, "") for c in sorted_codes)
        writer.writerow(row)

    output.seek(0)
    from flask import Response
    return Response(
        output.getvalue(),
        mimetype="text/csv; charset=utf-8",
        headers={"Content-Disposition": "attachment;filename=y4_students.csv"},
    )


# ---------------------------------------------------------------------------
# Prompt 迭代测试台 (Prompt Lab)
# ---------------------------------------------------------------------------
PROMPT_VERSIONS_DIR = BASE_DIR / "prompts" / "versions"
PROMPT_VERSIONS_DIR.mkdir(parents=True, exist_ok=True)


def _get_next_prompt_version() -> int:
    """Scan prompts/versions/ and return next version number."""
    existing = sorted(PROMPT_VERSIONS_DIR.glob("ai_interpreter_v*.md"))
    if not existing:
        return 1
    nums = []
    for f in existing:
        try:
            n = int(f.stem.split("_v")[1])
            nums.append(n)
        except (ValueError, IndexError):
            pass
    return max(nums) + 1 if nums else 1


@app.route("/prompt-lab")
@page_login_required
def prompt_lab_page():
    """Prompt iteration test bench."""
    return render_template("prompt_lab.html")


@app.route("/api/prompt-lab/run", methods=["POST"])
def prompt_lab_run():
    """Run AI interpretation using current prompt + report_data.json."""
    import urllib.request as _ureq
    import time as _time

    data = request.get_json(force=True)
    user_message = (data.get("message") or "请给出这份 Y4 报告的完整解读").strip()

    # 1) Read prompt
    prompt_path = BASE_DIR / "prompts" / "ai_interpreter.md"
    system_prompt = prompt_path.read_text(encoding="utf-8") if prompt_path.exists() else "你是测评解读助手。"

    # 2) Read report data
    report_path = DATA_DIR / "report_data.json"
    if not report_path.exists():
        return jsonify({"ok": False, "error": "没有测试数据 (report_data.json 不存在)"}), 400

    report_data = json.loads(report_path.read_text(encoding="utf-8"))
    schema_items = report_data.get("schema_124", [])
    data_text = "\n".join(
        f"{it.get('code','?')} {it.get('label','?')}：{it.get('value', '—')}"
        for it in schema_items if it.get("value")
    )
    student = report_data.get("student", {})
    student_text = f"学生：{student.get('name','—')}，{student.get('gender','—')}，{student.get('grade','—')}"
    context = f"{student_text}\n\n测评数据：\n{data_text}"

    # 3) Build messages
    messages = [
        {"role": "system", "content": system_prompt + "\n\n以下是学生测评数据：\n" + context},
        {"role": "user", "content": user_message},
    ]

    # 4) Call DashScope
    dashscope_key = os.environ.get("DASHSCOPE_API_KEY", extract.DEFAULT_DASHSCOPE_KEY).strip()
    url = "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions"
    payload = json.dumps({
        "model": os.environ.get("AI_TEXT_MODEL", "qwen-plus"),
        "messages": messages,
        "temperature": float(os.environ.get("AI_TEMPERATURE", "0.5")),
        "max_tokens": 8192,
    }).encode("utf-8")
    req = _ureq.Request(
        url, data=payload,
        headers={"Authorization": f"Bearer {dashscope_key}",
                 "Content-Type": "application/json"},
        method="POST",
    )
    t0 = _time.time()
    try:
        with _ureq.urlopen(req, timeout=120) as resp:
            result = json.loads(resp.read().decode("utf-8"))
        reply = result["choices"][0]["message"]["content"]
        tokens_used = result.get("usage", {}).get("total_tokens", 0)
        elapsed_ms = int((_time.time() - t0) * 1000)
        return jsonify({"ok": True, "reply": reply, "tokens": tokens_used, "time_ms": elapsed_ms})
    except Exception as exc:
        return jsonify({"ok": False, "error": f"AI 调用失败: {exc}"}), 500


@app.route("/api/prompt-lab/iterate", methods=["POST"])
def prompt_lab_iterate():
    """Auto-modify prompt based on user feedback, then return new prompt."""
    import urllib.request as _ureq

    data = request.get_json(force=True)
    feedback = (data.get("feedback") or "").strip()
    rating = data.get("rating", 3)
    last_output = (data.get("last_output") or "").strip()

    if not feedback:
        return jsonify({"ok": False, "error": "反馈不能为空"}), 400

    # 1) Read current prompt
    prompt_path = BASE_DIR / "prompts" / "ai_interpreter.md"
    current_prompt = prompt_path.read_text(encoding="utf-8") if prompt_path.exists() else ""

    # 2) Save current version
    version_num = _get_next_prompt_version()
    old_version_path = PROMPT_VERSIONS_DIR / f"ai_interpreter_v{version_num}.md"
    old_version_path.write_text(current_prompt, encoding="utf-8")

    # 3) Build meta-prompt for AI to improve the prompt
    meta_prompt = f"""你是一个 prompt 优化专家。你需要根据用户的反馈，修改 Y4 测评解读的 system prompt。

以下是当前用于 Y4 测评解读的 system prompt：
---
{current_prompt}
---

以下是用户对这个 prompt 生成输出的反馈：
评分：{rating}/5
反馈：{feedback}

上次 AI 的输出（供参考）：
---
{last_output[:3000]}
---

请根据用户反馈，修改上面的 system prompt。
要求：
- 只修改需要改进的部分，保留好的部分
- 输出完整的修改后的 prompt（不是 diff，不是解释，直接输出 prompt 全文）
- 保持 Y4 四维框架（心力/精力/学习力/生涯力）作为唯一语言体系
- 保持指标编号引用规范
- 不要加任何前缀说明或后缀解释"""

    messages = [{"role": "user", "content": meta_prompt}]

    # 4) Call DashScope to get improved prompt
    dashscope_key = os.environ.get("DASHSCOPE_API_KEY", extract.DEFAULT_DASHSCOPE_KEY).strip()
    url = "https://dashscope.aliyuncs.com/compatible-mode/v1/chat/completions"
    payload = json.dumps({
        "model": os.environ.get("AI_TEXT_MODEL", "qwen-plus"),
        "messages": messages,
        "temperature": 0.3,
        "max_tokens": 8192,
    }).encode("utf-8")
    req = _ureq.Request(
        url, data=payload,
        headers={"Authorization": f"Bearer {dashscope_key}",
                 "Content-Type": "application/json"},
        method="POST",
    )
    try:
        with _ureq.urlopen(req, timeout=120) as resp:
            result = json.loads(resp.read().decode("utf-8"))
        new_prompt = result["choices"][0]["message"]["content"].strip()

        # 5) Write new prompt
        prompt_path.write_text(new_prompt, encoding="utf-8")

        # 6) Simple diff: find first and last differing lines
        old_lines = current_prompt.split("\n")
        new_lines = new_prompt.split("\n")
        diff_parts = []
        max_compare = min(len(old_lines), len(new_lines))
        first_diff = None
        last_diff = None
        for i in range(max_compare):
            if old_lines[i] != new_lines[i]:
                if first_diff is None:
                    first_diff = i
                last_diff = i
        if first_diff is not None:
            start = max(0, first_diff - 2)
            end = min(max_compare, last_diff + 3)
            diff_parts.append(f"改动区域（第 {start+1}-{end} 行附近）：")
            for i in range(start, end):
                marker = "→" if old_lines[i] != new_lines[i] else " "
                if old_lines[i] != new_lines[i]:
                    diff_parts.append(f"  {marker} 旧: {old_lines[i][:80]}")
                    diff_parts.append(f"  {marker} 新: {new_lines[i][:80]}")
        if len(new_lines) != len(old_lines):
            diff_parts.append(f"\n行数变化：{len(old_lines)} → {len(new_lines)}")
        diff_summary = "\n".join(diff_parts) if diff_parts else "无明显差异"

        return jsonify({
            "ok": True,
            "old_prompt": current_prompt,
            "new_prompt": new_prompt,
            "diff": diff_summary,
            "version": version_num,
        })
    except Exception as exc:
        return jsonify({"ok": False, "error": f"Prompt 迭代失败: {exc}"}), 500


@app.route("/api/prompt-lab/save", methods=["POST"])
def prompt_lab_save():
    """Manually save edited prompt."""
    data = request.get_json(force=True)
    prompt_text = data.get("prompt", "").strip()
    if not prompt_text:
        return jsonify({"ok": False, "error": "Prompt 不能为空"}), 400

    # Save old version first
    prompt_path = BASE_DIR / "prompts" / "ai_interpreter.md"
    if prompt_path.exists():
        version_num = _get_next_prompt_version()
        old_version_path = PROMPT_VERSIONS_DIR / f"ai_interpreter_v{version_num}.md"
        old_version_path.write_text(prompt_path.read_text(encoding="utf-8"), encoding="utf-8")

    prompt_path.write_text(prompt_text, encoding="utf-8")
    return jsonify({"ok": True})


@app.route("/api/prompt-lab/prompt", methods=["GET"])
def prompt_lab_get_prompt():
    """Get current prompt text."""
    prompt_path = BASE_DIR / "prompts" / "ai_interpreter.md"
    content = prompt_path.read_text(encoding="utf-8") if prompt_path.exists() else ""
    return jsonify({"ok": True, "prompt": content})


if __name__ == "__main__:
    if len(sys.argv) > 1 and sys.argv[1] == "run":
        print("="*60)
        print("开始执行完整流程: extract → validate → generate")
        print("="*60)
        print()

        try:
            print("[1/3] 提取数据 (extract)...")
            rc = extract.main()
            if rc != 0:
                print(f"提取失败 (返回 {rc})")
                sys.exit(1)
            print("提取成功")
            print()

            print("[2/3] 校验数据 (validate)...")
            try:
                validate.main()
                print("校验成功")
            except Exception as exc:
                print(f"校验警告: {exc}")
            print()

            print("[3/3] 生成 PDF (generate)...")
            apply_report_data()
            _generate_module.main()
            print("生成成功")
            print()

            pdf_path = OUTPUT_DIR / "report.pdf"
            if pdf_path.exists():
                print(f"✅ PDF 已生成: {pdf_path}")
                print(f"   大小: {pdf_path.stat().st_size / 1024:.1f} KB")
            else:
                print("❌ PDF 生成失败")
                sys.exit(1)

        except Exception as exc:
            tb = traceback.format_exc()
            print(f"❌ 执行失败: {exc}")
            print(tb)
            sys.exit(1)

    else:
        port = int(os.environ.get("PORT", 8000))
        app.run(host="0.0.0.0", port=port, debug=False, threaded=True)
